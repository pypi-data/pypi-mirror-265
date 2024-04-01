#!/usr/bin/env python
# -*- coding: utf-8 -*-
# Author: Gary
# Datetime: 2023/1/10 9:46
# IDE: PyCharm
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.shared import RGBColor


def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def colored_cell(cell, value):
    """
        根据 value 设置cell的字体颜色。
        85 < value <= 90, RGB_YELLOW
        90 < value <= 95, RGB_ORANGE
        95 < value, RGB_RED
    """
    RGB_YELLOW = 255, 165, 0
    RGB_ORANGE = 255, 102, 0
    RGB_RED = 255, 0, 0
    if value >= 100:
        value = 99.99
    run = cell.paragraphs[0].add_run(str("%.2f" % value) + "%")
    if 85 < value <= 90:
        run.font.color.rgb = RGBColor(*RGB_YELLOW)
    elif 90 < value <= 95:
        run.font.color.rgb = RGBColor(*RGB_ORANGE)
    elif 95 < value:
        run.font.color.rgb = RGBColor(*RGB_RED)


class PyDocs:
    def __init__(self, doc):
        self._doc = doc

    def add_summarize_table(self, style="Medium Shading 1 Accent 1"):
        tbl = self._doc.add_table(rows=1, cols=12, style=style)
        tbl.autofit = False
        tbl.cell(0, 0).width, tbl.cell(0, 0).text = Inches(1), "主机组名称"
        tbl.cell(0, 1).width, tbl.cell(0, 1).text = Inches(0.5), "主机数量"
        tbl.cell(0, 2).width, tbl.cell(0, 2).text = Inches(0.75), "CPU平均利用率"
        tbl.cell(0, 3).width, tbl.cell(0, 3).text = Inches(0.75), "内存总量"
        tbl.cell(0, 4).width, tbl.cell(0, 4).text = Inches(0.75), "内存最高利用率"
        tbl.cell(0, 5).width, tbl.cell(0, 5).text = Inches(0.75), "内存最低利用率"
        tbl.cell(0, 6).width, tbl.cell(0, 6).text = Inches(0.75), "内存平均利用率"
        tbl.cell(0, 7).width, tbl.cell(0, 7).text = Inches(0.75), "磁盘总量"
        tbl.cell(0, 8).width, tbl.cell(0, 8).text = Inches(0.75), "磁盘最高使用率"
        tbl.cell(0, 9).width, tbl.cell(0, 9).text = Inches(0.75), "磁盘最低使用率"
        tbl.cell(0, 10).width, tbl.cell(0, 10).text = Inches(0.75), "磁盘平均使用率"
        tbl.cell(0, 11).width, tbl.cell(0, 11).text = Inches(0.75), "严重告警数量"
        return tbl

    def add_vm_table(self, style="Medium Shading 1 Accent 1"):
        tbl = self._doc.add_table(rows=1, cols=3, style=style)
        tbl.cell(0, 0).width, tbl.cell(0, 0).text = Inches(3.0), "主机组名称"
        tbl.cell(0, 1).width, tbl.cell(0, 1).text = Inches(4.0), "主机名称"
        tbl.cell(0, 2).width, tbl.cell(0, 2).text = Inches(0.7), "CPU平均使用率"
        return tbl

    def add_total_tbl(self, style="Medium Shading 1 Accent 1"):
        tbl = self._doc.add_table(rows=4, cols=2, style=style)
        tbl.cell(0, 0).text = "统计日期"
        tbl.cell(1, 0).text = "主机组数量"
        tbl.cell(2, 0).text = "主机数量"
        tbl.cell(3, 0).text = "严重告警数量"
        return tbl

    def add_detail_table(self, style="Medium Shading 1 Accent 1"):
        tbl = self._doc.add_table(rows=3, cols=6, style=style)
        cell_left = tbl.cell(0, 0)
        cell_right = tbl.cell(0, 5)
        header = cell_left.merge(cell_right)
        header.text = "各级别告警数量"
        tbl.cell(1, 0).text = "未分类"
        tbl.cell(1, 1).text = "通知"
        tbl.cell(1, 2).text = "警示"
        tbl.cell(1, 3).text = "严重"
        tbl.cell(1, 4).text = "危险"
        tbl.cell(1, 5).text = "灾难"
        return tbl

    def add_event_grp_tbl(self, style="Medium Shading 1 Accent 1"):
        tbl = self._doc.add_table(rows=1, cols=3, style=style)
        tbl.cell(0, 0).width = Inches(5)
        tbl.cell(0, 1).width = Inches(1)
        tbl.cell(0, 2).width = Inches(2.5)
        tbl.cell(0, 0).text = "主机组名称"
        tbl.cell(0, 1).text = "主机数量"
        tbl.cell(0, 2).text = "严重告警数量"
        return tbl

    def add_event_host_tbl(self, style="Medium Shading 1 Accent 1"):
        tbl = self._doc.add_table(rows=1, cols=3, style=style)
        tbl.cell(0, 0).width = Inches(3.0)
        tbl.cell(0, 1).width = Inches(3.4)
        tbl.cell(0, 2).width = Inches(1.3)
        tbl.cell(0, 0).text = "主机组名称"
        tbl.cell(0, 1).text = "主机名称"
        tbl.cell(0, 2).text = "严重告警数量"
        return tbl

    def add_mem_grp_tbl(self, style="Medium Shading 1 Accent 1"):
        tbl = self._doc.add_table(rows=1, cols=4, style=style)
        tbl.cell(0, 0).width = Inches(5)
        tbl.cell(0, 1).width = Inches(0.5)
        tbl.cell(0, 2).width = Inches(1.7)
        tbl.cell(0, 3).width = Inches(1.3)
        tbl.cell(0, 0).text = "主机组名称"
        tbl.cell(0, 1).text = "主机数量"
        tbl.cell(0, 2).text = "内存平均使用率"
        tbl.cell(0, 3).text = "内存总量"
        return tbl

    def add_mem_host_tbl(self, style="Medium Shading 1 Accent 1"):
        tbl = self._doc.add_table(rows=1, cols=4, style=style)
        tbl.cell(0, 0).width = Inches(3.0)
        tbl.cell(0, 1).width = Inches(3.4)
        tbl.cell(0, 2).width = Inches(1.3)
        tbl.cell(0, 3).width = Inches(1.3)
        tbl.cell(0, 0).text = "主机组名称"
        tbl.cell(0, 1).text = "主机名称"
        tbl.cell(0, 2).text = "内存平均使用率"
        tbl.cell(0, 3).text = "内存总量"
        return tbl

    def add_cpu_grp_tbl(self, style="Medium Shading 1 Accent 1"):
        tbl = self._doc.add_table(rows=1, cols=3, style=style)
        tbl.cell(0, 0).width = Inches(5)
        tbl.cell(0, 1).width = Inches(1.0)
        tbl.cell(0, 2).width = Inches(1.2)
        tbl.cell(0, 0).text = "主机组名称"
        tbl.cell(0, 1).text = "主机数量"
        tbl.cell(0, 2).text = "CPU平均使用率"
        return tbl

    def add_cpu_host_tbl(self, style="Medium Shading 1 Accent 1"):
        tbl = self._doc.add_table(rows=1, cols=3, style=style)
        tbl.cell(0, 0).width = Inches(3.5)
        tbl.cell(0, 1).width = Inches(3.5)
        tbl.cell(0, 2).width = Inches(0.7)
        tbl.cell(0, 0).text = "主机组名称"
        tbl.cell(0, 1).text = "主机名称"
        tbl.cell(0, 2).text = "CPU平均使用率"
        return tbl

    def add_disk_grp_tbl(self, style="Medium Shading 1 Accent 1"):
        tbl = self._doc.add_table(rows=1, cols=4, style=style)
        tbl.cell(0, 0).width = Inches(5)
        tbl.cell(0, 1).width = Inches(0.5)
        tbl.cell(0, 2).width = Inches(1.7)
        tbl.cell(0, 3).width = Inches(1.3)
        tbl.cell(0, 0).text = "主机组名称"
        tbl.cell(0, 1).text = "主机数量"
        tbl.cell(0, 2).text = "磁盘平均使用率"
        tbl.cell(0, 3).text = "磁盘总量"
        return tbl

    def add_disk_disk_tbl(self, style="Medium Shading 1 Accent 1"):
        tbl = self._doc.add_table(rows=1, cols=4, style=style)
        tbl.cell(0, 0).width = Inches(3.0)
        tbl.cell(0, 1).width = Inches(3.4)
        tbl.cell(0, 2).width = Inches(1.3)
        tbl.cell(0, 3).width = Inches(1.3)
        tbl.cell(0, 0).text = "主机组名称"
        tbl.cell(0, 1).text = "主机名称"
        tbl.cell(0, 2).text = "磁盘平均使用率"
        tbl.cell(0, 3).text = "磁盘总量"
        return tbl

    def add_para(self, run_, pt_, rgb_, alignment, para_content=""):
        para = self._doc.add_paragraph(para_content)
        para.alignment = alignment
        run = para.add_run(run_)
        run.bold = True
        run.font.size = Pt(pt_)
        run.font.color.rgb = RGBColor(*rgb_)

    def add_heading(self, level, run_, pt_, font_name="微软雅黑", qn_=qn("w:eastAsia"), heading=""):
        heading = self._doc.add_heading(heading, level=level).add_run(run_)
        heading.font.name = font_name
        heading.font.size = Pt(pt_)
        heading._element.rPr.rFonts.set(qn_, font_name)
