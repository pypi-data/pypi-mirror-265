#!/usr/bin/env python
# -*- coding: utf-8 -*-
# Author: Gary
# Datetime: 2022/12/28 17:33
# IDE: PyCharm
"""
    根据起止日期获取 Zabbix 资源使用率和排行榜报告。
"""
import argparse
import time
import logging
import math
import os
from collections import defaultdict, Counter
from itertools import groupby
from operator import itemgetter
from datetime import datetime
from dateutil.parser import parse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from lib.utils.zbxapis import ZabbixApiGet
from lib.utils.format import make_timestamp, convert_unit, convert_pdf, DiskCache
from lib.utils.format import get_value, jmes_search
from lib.utils.docs import *


class ZbxReport:
    def __init__(self, api, start, end):
        self._api = api
        self.start = start
        self.end = end
        self._cache = DiskCache()

    @property
    def server_hostid(self):
        """
            获取 "Zabbix Server" 主机的 hostid：
        :return:
        """
        if self._cache.get_cache("hostid_zbx_server"):
            return self._cache.get_cache("hostid_zbx_server")
        if not self._cache.get_cache("hostid_zbx_server"):
            server_host = self._api.get_hts(filter_={"host": "Zabbix server"})
            if server_host:
                self._cache.set_cache(
                    "hostid_zbx_server",
                    server_host[0].get("hostid")
                )
                return self._cache.get_cache("hostid_zbx_server")

    @property
    def htgrps(self):
        """
            获取 Zabbix 主机组信息：
        :return:
        """
        if self._cache.get_cache("htgrps_normal"):
            return list(self._cache.get_cache("htgrps_normal"))
        if not self._cache.get_cache("htgrps_normal"):
            htgrps = self._api.get_ht_grps(
                output=["name"],
                selecthosts=["hostid", "name"],
                real_hosts=True,
                with_monitored_items=True,
                filter_={"flags": 0}
            )
            self._cache.set_cache("htgrps_normal", htgrps, expire=300)
            return list(self._cache.get_cache("htgrps_normal"))

    @property
    def vm_hosts(self):
        """
            获取 Host Inventory Type 为 "VM" 的主机信息：
        :return:
        """
        hosts = self._api.get_hts(
            with_monitored_items=True,
            searchinventory={"type": "VM"},
            filter_={"flags": 0},
            output=["hostid"],
            selectinterfaces=["available"]
        )
        hosts = jmes_search(
            jmes_rexp=get_value(
                section="JMES",
                option="SEARCH_AVAILABLE_HOSTS"
            ),
            data=hosts
        )
        return hosts

    def get_host_items(self, group: dict):
        """
            获取 Zabbix 主机指定项的 Item 信息：
        :param group:
        :return:
        """
        if self._cache.get_cache("items_" + str(group.get("groupid"))):
            return list(self._cache.get_cache("items_" + str(group.get("groupid"))))
        if not self._cache.get_cache("items_" + str(group.get("groupid"))):
            items = self._api.get_items(
                hostids=[host.get("hostid") for host in group.get("hosts")],
                output=["name", "key_", "hostid"],
                monitored=True,
                filter_={
                    "key_": [
                        "vfs.fs.totalsize",
                        "vfs.fs.usedsize",
                        "system.cpu.util[,idle]",
                        "vm.memory.size[used]",
                        "vm.memory.size[total]"
                    ],
                    "state": 0
                }
            )
            self._cache.set_cache(
                "items_" + str(group.get("groupid")),
                items,
                expire=60
            )
            return list(self._cache.get_cache("items_" + str(group.get("groupid"))))

    @property
    def items(self):
        """
            获取 "Zabbix Server" 主机为开启状态的 Items 信息：
        :return:
        """
        if self._cache.get_cache("zbx_server_items"):
            return list(self._cache.get_cache("zbx_server_items"))
        if not self._cache.get_cache("zbx_server_items"):
            items = self._api.get_items(
                hostids=self.server_hostid,
                output=["name", "key_"],
                monitored=True,
                filter_={"state": 0}
            )
            self._cache.set_cache(
                "zbx_server_items",
                items,
                expire=60
            )
            return list(self._cache.get_cache("zbx_server_items"))

    def get_itemids(self, name: str):
        """
            获取 Zabbix Items 的 id：
        :param name:
        :return:
        """
        if self._cache.get_cache("itemids_" + name):
            return list(self._cache.get_cache("itemids_" + name))
        if not self._cache.get_cache("itemids_" + name):
            itemids = [item.get("itemid") for item in self.items if item.get("name") == name]
            self._cache.set_cache(
                "itemids_" + name,
                itemids,
                expire=60
            )
            return list(self._cache.get_cache("itemids_" + name))

    def getcalc(self, itemids):
        """
            获取【计算型】监控项指定时间范围内的最大值、最小值、平均值：
        :param itemids:
        :return:
        """
        trends = self._api.get_trends(
            itemids=itemids,
            time_from=make_timestamp(self.start),
            time_till=make_timestamp(self.end)
        )
        if len(trends) != 0:
            values_min = []
            values_max = []
            values_avg = []
            for trend in trends:
                values_min.append(float(trend["value_min"]))
                values_max.append(float(trend["value_max"]))
                values_avg.append(float(trend["value_avg"]))
            avg_value = round(sum(values_avg) / len(values_avg), 2)
            min_value = min(values_min)
            max_value = max(values_max)
            return min_value, max_value, avg_value
        return 0, 0, 0

    def get_zbx_events(self, severities: list):
        """
            获取 Zabbix 监控主机的告警事件：
        :param severities:
        :return:
        """
        return self._api.get_events(
            countoutput=True,
            value=1,
            severities=severities,
            time_from=make_timestamp(self.start),
            time_till=make_timestamp(self.end)
        )


def get_word(api, path, start, end, topnum):
    """" 生成word统计报表 """
    document = Document()
    docs = PyDocs(document)
    zbx = ZbxReport(api, start, end)
    document.styles["Normal"].font.name = "微软雅黑"
    document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    docs.add_para(
        run_="\n鑫运运管平台\n监控统计分析月报\n\n",
        pt_=36,
        rgb_=(79, 129, 189),
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    docs.add_para(
        run_="\n\n" + end[0:4] + "年" + end[4:6] + "月",
        pt_=18,
        rgb_=(79, 129, 189),
        alignment=None
    )
    document.add_page_break()
    # 1.汇总信息页
    docs.add_heading(level=1, run_="一、汇总信息", pt_=20)
    # 1.1表头
    table_total = docs.add_total_tbl()
    # 1.2表内数据
    table_total.cell(0, 1).text = "{} - {}".format(
        time.strftime("%Y/%m/%d", time.strptime(start, "%Y%m%d")),
        time.strftime("%Y/%m/%d", time.strptime(end, "%Y%m%d")))
    # 获取主机组
    host_groups = zbx.htgrps
    # 主机组总数量
    groups_num = len(host_groups)
    # 主机总数量
    hosts_sum = []
    for grp in host_groups:
        hosts_sum += [host.get("hostid") for host in grp.get("hosts")]
    # 获取严重告警数量
    event_sum_num = zbx.get_zbx_events([3, 4, 5])
    table_total.cell(1, 1).text = str(groups_num)
    table_total.cell(2, 1).text = str(len(set(hosts_sum)))
    table_total.cell(3, 1).text = str(event_sum_num)
    run_event_number = document.add_paragraph("")
    run_event_number.paragraph_format.space_before = 15
    table_detail_number = docs.add_detail_table()
    # 获取对应告警级别数量
    for severity in range(6):
        event_num = zbx.get_zbx_events([severity])
        table_detail_number.cell(2, severity).text = str(event_num)
    docs.add_para(
        run_="注: `严重`、`危险`、`灾难` 三个等级的告警纳入严重告警统计",
        pt_=10,
        rgb_=(0, 139, 0),
        alignment=None
    )
    # 严重告警数量表格
    document.add_page_break()
    docs.add_heading(level=1, run_="二、严重告警数量排行", pt_=20)
    # 新增2级标题
    docs.add_heading(level=2, run_=f"1、严重告警数量最多的{topnum}个主机组", pt_=16)
    # 插入表格
    event_table_desc_group = docs.add_event_grp_tbl()
    document.add_page_break()
    # 严重告警数量最多的主机
    logging.info("\033[32m严重告警数量排行, 主机维度\033[0m")
    docs.add_heading(level=2, run_=f"2、严重告警数量最多的{topnum}台主机", pt_=16)
    event_table_desc_host = docs.add_event_host_tbl()
    # 2.详细统计信息页
    # 2.1 表头
    document.add_page_break()
    docs.add_heading(level=1, run_="三、主机组资源利用率概览", pt_=20)
    # 2.3.2 获取 zabbix server 中各 hostgroup的聚合 item
    # 根据 Average cpu utilization 监控项确定主机组
    valid_hostgroup_names = [
        item.get("name").split("group")[1].strip()
        for item in zbx.items
        if item.get("name").startswith("Average cpu utilization in group")
    ]
    host_groups = [g for g in host_groups if g.get("name") in valid_hostgroup_names]
    # 按主机数量排序主机组
    host_groups.sort(key=lambda x: len(x["hosts"]), reverse=True)
    # 2.3.3 设置按主机组维度统计数据的变量
    # 主机组维度按内存使用率avg数组
    memory_top_group = []
    # cpu利用率数组（主机组维度）
    cpu_top_group = []
    # 磁盘使用率（主机组维度）
    filesystem_top_group = []
    # 告警数量 (主机组维度)
    event_count_group = []
    # 主机维度按内存使用率avg数组
    memory_top_host = []
    # cpu利用率数组（主机维度）
    cpu_top_host = []
    # 磁盘使用率（主机维度）
    filesystem_top_host = []
    # 告警数量 (主机组维度)
    event_count_host = []
    # 2.3.4 填充表格数据
    summarize_row_count = 0
    for index, group in enumerate(host_groups):
        group_name = group.get("name")
        logging.info("\033[33m正在处理数据……主机组：%s\033[0m", group_name)
        logging.info("\033[33m开始时间：%s\033[0m", str(datetime.now()))
        summarize_table = None
        if summarize_row_count == 0:
            summarize_table = docs.add_summarize_table()
        host_num = len(group.get("hosts"))
        row = summarize_table.add_row()
        row.cells[0].text = group_name
        row.cells[1].text = str(host_num)
        # group_name 5个字一行, 计算共占多少行
        summarize_row_count += math.ceil(len(group_name) / 5)
        # 获取cpu利用率
        _, _, avg_v = zbx.getcalc(zbx.get_itemids(f"Average cpu utilization in group {group_name}"))
        colored_cell(row.cells[2], avg_v)
        # 保留信息
        cpu_top_group.append(
            {
                "groupname": group_name,
                "hostnum": host_num,
                "cpu_utilization": avg_v
            }
        )
        # 获取内存总量
        _, _, avg_v = zbx.getcalc(zbx.get_itemids(f"Total memory in group {group_name}"))
        row.cells[3].text = convert_unit(avg_v)
        memory_dic = {
            "groupname": group_name,
            "hostnum": host_num,
            "memory_total": avg_v
        }
        # 获取内存利用率
        min_v, max_v, avg_v = zbx.getcalc(
            zbx.get_itemids(f"Memory utilization in group {group_name}")
        )
        colored_cell(row.cells[4], max_v)
        colored_cell(row.cells[5], min_v)
        colored_cell(row.cells[6], avg_v)
        memory_dic["memory_utilization"] = avg_v
        memory_top_group.append(memory_dic)
        # 获取磁盘总量
        _, _, avg_v = zbx.getcalc(zbx.get_itemids(f"Total disk space in {group_name}"))
        row.cells[7].text = convert_unit(avg_v)
        filesystem_dic = {
            "groupname": group_name,
            "hostnum": host_num,
            "filesystem_total": avg_v
        }
        # 获取磁盘使用率
        min_v, max_v, avg_v = zbx.getcalc(
            zbx.get_itemids(f"Used disk space in {group_name} (percentage)")
        )
        colored_cell(row.cells[8], max_v)
        colored_cell(row.cells[9], min_v)
        colored_cell(row.cells[10], avg_v)
        filesystem_dic["filesystem_utilization"] = avg_v
        filesystem_top_group.append(filesystem_dic)
        # 按主机维度处理信息，包括过滤警告，以及获取对应主机的分析数据
        group_host_keys = defaultdict(dict)
        for host_item in zbx.get_host_items(group=group):
            host_name = [
                host["name"] for host in group["hosts"]
                if host["hostid"] == host_item["hostid"]
            ][0]
            group_host_keys[host_name][host_item["key_"]] = host_item["itemid"]
            group_host_keys[host_name]["hostid"] = host_item["hostid"]
        for host_name, host_keys in group_host_keys.items():
            # 获取主机分析数据
            # 内存 used 、 total
            if host_keys.get("vm.memory.size[total]"):
                _, _, mem_avg_used = zbx.getcalc(host_keys["vm.memory.size[used]"])
                _, _, mem_avg_total = zbx.getcalc(host_keys["vm.memory.size[total]"])
                if mem_avg_total != 0:
                    # 内存使用率
                    memory_top_host.append(
                        {
                            "hostname": host_name,
                            "memory_utilization": 100 * mem_avg_used / mem_avg_total,
                            "memory_total": mem_avg_total,
                            "groupname": group_name
                        }
                    )
            # cpu 使用率
            if host_keys.get("system.cpu.util[,idle]"):
                _, _, cpu_avg_idle = zbx.getcalc(host_keys["system.cpu.util[,idle]"])
                if cpu_avg_idle != 0:
                    cpu_top_host.append(
                        {
                            "hostname": host_name,
                            "hostid": host_keys.get("hostid"),
                            "cpu_utilization": 100 - cpu_avg_idle,
                            "groupname": group_name
                        }
                    )
            # 磁盘 used 、 total
            if host_keys.get("vfs.fs.totalsize") and host_keys.get("vfs.fs.usedsize"):
                _, _, disk_avg_used = zbx.getcalc(host_keys.get("vfs.fs.usedsize"))
                _, _, disk_avg_total = zbx.getcalc(host_keys.get("vfs.fs.totalsize"))
                # 磁盘 使用率
                if disk_avg_used != 0:
                    filesystem_top_host.append(
                        {
                            "hostname": host_name,
                            "filesystem_utilization": 100 * disk_avg_used / disk_avg_total,
                            "filesystem_total": disk_avg_total,
                            "groupname": group_name
                        }
                    )
        events = api.get_events(
            output=["eventid"],
            selecthosts=["name"],
            hostids=[host.get("hostid") for host in group.get("hosts")],
            value=1,
            severities=[3, 4, 5],
            time_from=make_timestamp(start),
            time_till=make_timestamp(end)
        )
        row.cells[11].text = str(len(events))
        # 主机组维度 严重告警事件数量
        event_count_dic = {
            "groupname": group_name,
            "hostnum": host_num,
            "events_count": len(events)
        }
        event_count_group.append(event_count_dic)
        # 主机维度 严重告警事件数量
        events_by_host = Counter(e['hosts'][0]['name'] for e in events if e['hosts'])
        for host_name in events_by_host:
            event_count_host.append(
                {
                    "hostname": host_name,
                    "events_count": events_by_host[host_name],
                    "groupname": group_name
                }
            )
        if index == len(host_groups) - 1:
            document.add_page_break()
        elif summarize_row_count >= 18:
            summarize_row_count = 0
            document.add_page_break()
    # 主机组按严重告警数量排序desc
    event_count_group.sort(key=lambda x: x["events_count"], reverse=True)
    for i in range(min(topnum, len(event_count_group))):
        row = event_table_desc_group.add_row()
        row.cells[0].text = event_count_group[i]["groupname"]
        row.cells[1].text = str(event_count_group[i]["hostnum"])
        row.cells[2].text = str(event_count_group[i]["events_count"])
    event_count_host.sort(key=lambda x: x["events_count"], reverse=True)
    for i in range(min(topnum, len(event_count_host))):
        row = event_table_desc_host.add_row()
        row.cells[0].text = event_count_host[i]["groupname"]
        row.cells[1].text = event_count_host[i]["hostname"]
        row.cells[2].text = str(event_count_host[i]["events_count"])
    # 3. 内存使用率排行
    docs.add_heading(level=1, run_="四、内存使用率排行", pt_=20)
    docs.add_heading(level=2, run_=f"1、内存使用率最高的{topnum}个主机组", pt_=16)
    # 插入表格 按内存使用率排序desc
    table_memory_group_desc = docs.add_mem_grp_tbl()
    memory_top_group.sort(key=lambda x: x["memory_utilization"], reverse=True)
    for i in range(min(topnum, len(memory_top_group))):
        row = table_memory_group_desc.add_row()
        row.cells[0].text = memory_top_group[i]["groupname"]
        row.cells[1].text = str(memory_top_group[i]["hostnum"])
        colored_cell(row.cells[2], memory_top_group[i]["memory_utilization"])
        row.cells[3].text = str(convert_unit(memory_top_group[i]["memory_total"]))
    document.add_page_break()
    docs.add_heading(level=2, run_=f"2、内存使用率最低的{topnum}个主机组", pt_=16)
    # 插入表格 按内存使用率排序asc
    table_memory_group_asc = docs.add_mem_grp_tbl()
    memory_top_group.sort(key=lambda x: x["memory_utilization"])
    for i in range(min(topnum, len(memory_top_group))):
        row = table_memory_group_asc.add_row()
        row.cells[0].text = memory_top_group[i]["groupname"]
        row.cells[1].text = str(memory_top_group[i]["hostnum"])
        colored_cell(row.cells[2], memory_top_group[i]["memory_utilization"])
        row.cells[3].text = str(convert_unit(memory_top_group[i]["memory_total"]))
    document.add_page_break()
    docs.add_heading(level=2, run_=f"3、内存使用率最高的{topnum}台主机", pt_=16)
    # 插入表格
    table_memory_host_desc = docs.add_mem_host_tbl()
    memory_top_host.sort(key=itemgetter("memory_utilization"))
    memory_top_host_groupby = []
    for hostname, hosts_iter in groupby(memory_top_host, key=itemgetter("hostname")):
        hosts = list(hosts_iter)
        memory_top_host_groupby.append(
            {
                "hostname": hostname,
                "memory_utilization": hosts[0]["memory_utilization"],
                "memory_total": hosts[0]["memory_total"],
                "groupname": ','.join(h['groupname'] for h in hosts)
            }
        )
    memory_top_host_groupby.sort(key=itemgetter("memory_utilization"), reverse=True)
    for i in range(min(topnum, len(memory_top_host))):
        row = table_memory_host_desc.add_row()
        row.cells[0].text = memory_top_host_groupby[i]["groupname"]
        row.cells[1].text = memory_top_host_groupby[i]["hostname"]
        colored_cell(row.cells[2], memory_top_host_groupby[i]["memory_utilization"])
        row.cells[3].text = str(convert_unit(memory_top_host_groupby[i]["memory_total"]))
    document.add_page_break()
    docs.add_heading(level=2, run_=f"4、内存使用率最低的{topnum}台主机", pt_=16)
    # 插入表格
    table_memory_host_asc = docs.add_mem_host_tbl()
    memory_top_host_groupby.sort(key=itemgetter("memory_utilization"))
    for i in range(min(topnum, len(memory_top_host))):
        row = table_memory_host_asc.add_row()
        row.cells[0].text = memory_top_host_groupby[i]["groupname"]
        row.cells[1].text = memory_top_host_groupby[i]["hostname"]
        colored_cell(row.cells[2], memory_top_host_groupby[i]["memory_utilization"])
        row.cells[3].text = str(convert_unit(memory_top_host_groupby[i]["memory_total"]))
    document.add_page_break()
    docs.add_heading(level=1, run_="五、CPU使用率排行", pt_=20)
    docs.add_heading(level=2, run_=f"1、CPU使用率最高的{topnum}个主机组", pt_=16)
    # 插入表格 按cpu使用率排序主机组维度 desc
    table_cpu_group_desc = docs.add_cpu_grp_tbl()
    cpu_top_group.sort(key=lambda x: x["cpu_utilization"], reverse=True)
    for i in range(min(topnum, len(cpu_top_group))):
        row = table_cpu_group_desc.add_row()
        row.cells[0].text = cpu_top_group[i]["groupname"]
        row.cells[1].text = str(cpu_top_group[i]["hostnum"])
        colored_cell(row.cells[2], cpu_top_group[i]["cpu_utilization"])
    document.add_page_break()
    docs.add_heading(level=2, run_=f"2、CPU使用率最低的{topnum}个主机组", pt_=16)
    # 插入表格 按cpu使用率排序 主机组维度 asc
    table_cpu_group_asc = docs.add_cpu_grp_tbl()
    cpu_top_group.sort(key=lambda x: x["cpu_utilization"])
    for i in range(min(topnum, len(cpu_top_group))):
        row = table_cpu_group_asc.add_row()
        row.cells[0].text = cpu_top_group[i]["groupname"]
        row.cells[1].text = str(cpu_top_group[i]["hostnum"])
        colored_cell(row.cells[2], cpu_top_group[i]["cpu_utilization"])
    document.add_page_break()
    docs.add_heading(level=2, run_=f"3、CPU使用率最高的{topnum}台主机", pt_=16)
    # 插入表格 cpu使用率主机维度 desc
    table_cpu_host_desc = docs.add_cpu_host_tbl()
    cpu_top_host.sort(key=itemgetter("cpu_utilization"))
    cpu_top_host_groupby = []
    for hostname, hosts_iter in groupby(cpu_top_host, key=itemgetter("hostname")):
        hosts = list(hosts_iter)
        cpu_top_host_groupby.append(
            {
                "hostname": hostname,
                "cpu_utilization": hosts[0]["cpu_utilization"],
                "groupname": ','.join(h['groupname'] for h in hosts)
            }
        )
    cpu_top_host_groupby.sort(key=itemgetter("cpu_utilization"), reverse=True)
    for i in range(min(topnum, len(cpu_top_host_groupby))):
        row = table_cpu_host_desc.add_row()
        row.cells[0].text = cpu_top_host_groupby[i]["groupname"]
        row.cells[1].text = cpu_top_host_groupby[i]["hostname"]
        colored_cell(row.cells[2], cpu_top_host_groupby[i]["cpu_utilization"])
    document.add_page_break()
    docs.add_heading(level=2, run_=f"4、CPU使用率最低的{topnum}台主机", pt_=16)
    # 插入表格
    table_cpu_host_asc = docs.add_cpu_host_tbl()
    cpu_top_host_groupby.sort(key=itemgetter("cpu_utilization"))
    for i in range(min(topnum, len(cpu_top_host_groupby))):
        row = table_cpu_host_asc.add_row()
        row.cells[0].text = cpu_top_host_groupby[i]["groupname"]
        row.cells[1].text = cpu_top_host_groupby[i]["hostname"]
        colored_cell(row.cells[2], cpu_top_host_groupby[i]["cpu_utilization"])
    document.add_page_break()
    docs.add_heading(level=1, run_="六、磁盘使用率排行", pt_=20)
    docs.add_heading(level=2, run_=f"1、磁盘使用率最高的{topnum}个主机组", pt_=16)
    # 插入表格主机组按磁盘使用率排序desc
    table_disk_group_desc = docs.add_disk_grp_tbl()
    filesystem_top_group.sort(key=lambda x: x["filesystem_utilization"], reverse=True)
    for i in range(min(topnum, len(filesystem_top_group))):
        row = table_disk_group_desc.add_row()
        row.cells[0].text = filesystem_top_group[i]["groupname"]
        row.cells[1].text = str(filesystem_top_group[i]["hostnum"])
        colored_cell(row.cells[2], filesystem_top_group[i]["filesystem_utilization"])
        row.cells[3].text = str(convert_unit(
            filesystem_top_group[i]["filesystem_total"]))
    document.add_page_break()
    docs.add_heading(level=2, run_=f"2、磁盘使用率最低的{topnum}个主机组", pt_=16)
    # 插入表格  主机组按磁盘使用率排序asc
    table_disk_group_asc = docs.add_disk_grp_tbl()
    filesystem_top_group.sort(key=lambda x: x["filesystem_utilization"])
    for i in range(min(topnum, len(filesystem_top_group))):
        row = table_disk_group_asc.add_row()
        row.cells[0].text = filesystem_top_group[i]["groupname"]
        row.cells[1].text = str(filesystem_top_group[i]["hostnum"])
        colored_cell(row.cells[2], filesystem_top_group[i]["filesystem_utilization"])
        row.cells[3].text = str(convert_unit(filesystem_top_group[i]["filesystem_total"]))
    document.add_page_break()
    docs.add_heading(level=2, run_=f"3、磁盘使用率最高的{topnum}台主机", pt_=16)
    # 插入表格 磁盘使用率 主机维度 desc
    table_disk_disk_desc = docs.add_disk_disk_tbl()
    filesystem_top_host.sort(key=itemgetter("hostname"))
    filesystem_top_host_groupby = []
    for hostname, hosts_iter in groupby(filesystem_top_host, key=itemgetter("hostname")):
        hosts = list(hosts_iter)
        filesystem_top_host_groupby.append(
            {
                "hostname": hostname,
                "filesystem_utilization": hosts[0]["filesystem_utilization"],
                "filesystem_total": hosts[0]["filesystem_total"],
                "groupname": ','.join(h['groupname'] for h in hosts)
            }
        )
    filesystem_top_host_groupby.sort(
        key=itemgetter("filesystem_utilization"), reverse=True)
    for i in range(min(topnum, len(filesystem_top_host_groupby))):
        row = table_disk_disk_desc.add_row()
        row.cells[0].text = filesystem_top_host_groupby[i]["groupname"]
        row.cells[1].text = filesystem_top_host_groupby[i]["hostname"]
        colored_cell(row.cells[2], filesystem_top_host_groupby[i]["filesystem_utilization"])
        row.cells[3].text = str(convert_unit(filesystem_top_host_groupby[i]["filesystem_total"]))
    document.add_page_break()
    docs.add_heading(level=2, run_=f"4、磁盘使用率最低的{topnum}台主机", pt_=16)
    # 插入表格 磁盘使用率 主机维度 asc
    table_disk_disk_asc = docs.add_disk_disk_tbl()
    filesystem_top_host_groupby.sort(key=itemgetter("filesystem_utilization"))
    for i in range(min(topnum, len(filesystem_top_host_groupby))):
        row = table_disk_disk_asc.add_row()
        row.cells[0].text = filesystem_top_host_groupby[i]["groupname"]
        row.cells[1].text = filesystem_top_host_groupby[i]["hostname"]
        colored_cell(row.cells[2], filesystem_top_host_groupby[i]["filesystem_utilization"])
        row.cells[3].text = str(convert_unit(filesystem_top_host_groupby[i]["filesystem_total"]))
    # cpu使用率低于1%的虚拟机列表
    document.add_page_break()
    docs.add_heading(level=1, run_="七、CPU使用率较低的虚拟机", pt_=20)
    docs.add_heading(level=2, run_="1、CPU使用率低于1%的虚拟机", pt_=16)
    # 插入表格
    vm_page_num = 1
    vm_page_row_count = 0
    vm_nodata = True
    for vm in zbx.vm_hosts:
        vm_cpu_info = [host for host in cpu_top_host if host["hostid"] == vm["hostid"]]
        if not vm_cpu_info:
            continue
        if vm_cpu_info[0].get("cpu_utilization", 0) < 1:
            if vm_page_row_count == 0:
                table_vm_cpu = docs.add_vm_table()
                row = table_vm_cpu.add_row()
                row.cells[0].text = vm_cpu_info[0]["groupname"]
                row.cells[1].text = vm_cpu_info[0]["hostname"]
                colored_cell(row.cells[2], vm_cpu_info[0]["cpu_utilization"])
                vm_page_row_count += 1
                if (vm_page_num == 1 and vm_page_row_count >= 17) or (vm_page_row_count >= 21):
                    # 第一页满17行换页 非第一页满21行换页
                    vm_page_num += 1
                    vm_page_row_count = 0
                    vm_nodata = False
                    document.add_page_break()
    # 无数据则填充一行`无`
    if vm_nodata:
        table_vm_cpu = docs.add_vm_table()
        row = table_vm_cpu.add_row()
        for i in range(len(table_vm_cpu.columns)):
            row.cells[i].text = "无"
    # 设置纸张方向为横向
    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    # 遍历所有表格, 为每个单元格添加边框, 设置文字居中
    for _, table in enumerate(document.tables):
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                if c == 0:
                    continue  # 跳过第一列
                if r == 0:
                    # 表头用浅色边框
                    color = "#DDDDDD"
                else:
                    # 深色边框
                    color = "#7BA0CD"
                set_cell_border(
                    cell,
                    start={"sz": 1, "color": color, "val": "single"}
                )
                # 除第一列外，表格数据居中显示
                if c > 0:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 调整各级别告警数量表格样式
    for row in range(1, 3):
        for col in range(6):
            table_detail_number.cell(row, col).paragraphs[0].runs[0].font.size = Pt(12)
            table_detail_number.cell(row, col).paragraphs[0].runs[0].bold = False
            table_detail_number.cell(row, col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_detail_number.cell(
                row, col
            ).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(path)


def main(args):
    """main function"""
    zapi = ZabbixApiGet(args.zapi)
    start = datetime.strftime(parse(args.start), "%Y%m%d")
    output_docx = args.output or f"运管平台统计分析月报{start[:4]}年{start[4:6]}月.docx"
    get_word(
        api=zapi,
        path=output_docx,
        start=start,
        end=datetime.strftime(parse(args.end), "%Y%m%d"),
        topnum=args.topnum
    )
    logging.info(
        "\033[32mWord报表导出完成：%s\033[0m",
        os.path.abspath(output_docx) if os.path.exists(output_docx) else ""
    )
    convert_pdf(api_secret=args.api_secret, output_docx=output_docx)


parser = argparse.ArgumentParser()
parser.add_argument(
    "-s",
    "--start",
    required=True,
    help="The Start Date",
    type=str
)
parser.add_argument(
    "-e",
    "--end",
    required=True,
    help="The End Date",
    type=str
)
parser.add_argument(
    "-o",
    "--output",
    type=str,
    help="output filename"
)
parser.add_argument(
    "-t",
    "--topnum",
    type=int,
    default=10
)
parser.add_argument(
    "-a",
    "--api-secret",
    type=str,
    help="convertapi api_secret"
)
parser.set_defaults(handler=main)
