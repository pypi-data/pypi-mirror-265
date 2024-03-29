import os
import mmap
import imghdr
import xlrd
import pandas as pd
import openpyxl
import docx
import json
# import fitz
import tkinter as tk
from tkinter import filedialog
# from win32com import client as wc
import PyPDF2
import configparser
from bs4 import BeautifulSoup
from lxml import etree
import execjs
import cv2
from moviepy.editor import VideoFileClip
from jinja2 import Environment, FileSystemLoader
from loguru import logger


class TxtFileHandle:
    @staticmethod
    def read_txt(file_path, return_list=True):
        with open(file_path, "r", encoding="utf-8") as f:
            lines = [line.strip("\n") for line in f]
        if return_list:
            return lines
        yield from lines  # 节省内存

    @staticmethod
    def write_txt(file_path, data, is_add=True):
        mode = "a" if is_add else "w"
        with open(file_path, mode, encoding="utf-8") as f:
            if isinstance(data, list):
                for info in data:
                    TxtFileHandle.save_data(f, info)
            else:
                TxtFileHandle.save_data(f, data)

    @staticmethod
    def save_data(f, info):
        f.write("%s\n" % info)
        f.flush()

    @staticmethod
    def get_file_count(file_path):
        count = 0
        with open(file_path, "r", encoding="utf-8") as f:
            with mmap.mmap(f.fileno(), 0, access=mmap.ACCESS_READ) as mm:
                while True:
                    line = mm.readline()
                    if not line:
                        break
                    count += 1
        return count


class XlsFileHandle:
    """xls文件相对于xlsx较老，不建议使用，若有特殊要求可以使用这个"""

    @staticmethod
    def read_xls(file_path, sheet_index=0, sheet_name=None):
        workbook = xlrd.open_workbook(file_path)
        if sheet_name is not None:
            sheet = workbook.sheet_by_name(sheet_name)
        else:
            sheet = workbook.sheet_by_index(sheet_index)

        for row in range(sheet.nrows):
            data = sheet.row_values(row)
            yield data

    @staticmethod
    def write_xls(dataframe, file_path, columns_name=None, header=True):
        """
        dataframe:[[]]DataFrame形式，列名columns_name = ['a','b',...]
        暂时没写追加方法
        """
        pd_data4 = pd.DataFrame(dataframe, columns=columns_name)
        pd_data4.to_excel(file_path, header=header, index=False)  # 加入header=False即不要列名


class XlsxFileHandle:
    @staticmethod
    def read_xlsx(file_path, sheet_name=None):
        workbook = openpyxl.load_workbook(file_path)
        if sheet_name:
            sheet = workbook[sheet_name]
        else:
            sheet = workbook.active
        for row in sheet.iter_rows(values_only=True):
            yield row
        workbook.close()

    @staticmethod
    def write_xlsx(file_path, dataframe, sheet_name="", columns_name=None, is_add=False):
        """
        columns_name:["a", "b", "c"]
        dataframe:[[]]
        """
        OtherFileHandle.create_file(file_path)
        if not is_add:
            workbook = openpyxl.Workbook()
        else:
            workbook = openpyxl.load_workbook(file_path)
        if sheet_name:
            sheet = workbook.create_sheet(sheet_name)
        else:
            sheet = workbook.active
        if columns_name is not None:
            sheet.append(columns_name)
        for data in dataframe:
            sheet.append(data)
        workbook.save(file_path)
        workbook.close()


# class DocFileHandle:
#     """doc文档在python中一般是转换成docx文档进行处理"""
#
#     def __init__(self):
#         self.word = wc.Dispatch("Word.Application")  # 打开word应用程序
#
#     def doc_to_docx(self, doc_file_path, docx_file_path):
#         doc = self.word.Documents.Open(doc_file_path)  # 打开word文件
#         doc.SaveAs(docx_file_path, 12)  # 另存为后缀为".docx"的文件，其中参数12指docx文件
#         doc.Close()  # 关闭原来word文件
#
#     def quit(self):
#         self.word.Quit()


class DocxFileHandle:
    @staticmethod
    def read_docx(file_path, need_tables=False):
        text = ""
        file = docx.Document(file_path)

        # 获取页眉内容
        for section in file.sections:
            header = section.header
            if header is not None:
                for paragraph in header.paragraphs:
                    text += paragraph.text

        # 获取页脚内容
        for section in file.sections:
            footer = section.footer
            if footer is not None:
                for paragraph in footer.paragraphs:
                    text += paragraph.text

        # 获取主体内容
        for para in file.paragraphs:
            text += para.text
        if not need_tables:
            return text

        # 获取表格文本数据
        tables_text_list = []
        for table in file.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text not in tables_text_list:
                        tables_text_list.append(cell.text.strip())
        s = "".join(tables_text_list)
        text += s
        return text


class JsonFileHandle:
    @staticmethod
    def read_json(file_path):
        with open(file_path, "r", encoding="utf-8") as f:
            file = json.load(f)
        return file

    @staticmethod
    def write_json(new_file, file_path):
        """indent参数表示每个级别缩进的空格数"""
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(new_file, f, indent=4, ensure_ascii=False)


class IniFileHandle:
    def __init__(self, config_file_path):
        self.config_file = config_file_path
        self.config = configparser.ConfigParser()
        self.config.read(self.config_file, encoding="utf-8")

    def get_value(self, section, key):
        try:
            return json.loads(self.config.get(section, key))
        except (configparser.NoSectionError, configparser.NoOptionError):
            logger.error("not found value and section is %s key is %s set default is None" % (section, key))
            return None

    def set_value(self, section, key, _value):
        if section not in self.config:
            self.config.add_section(section)
        self.config.set(section, key, _value)

    def save_config(self):
        with open(self.config_file, 'w') as configfile:
            self.config.write(configfile)


class PdfFileHandle:
    @staticmethod
    def read_pdf(file_path):
        with open(file_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            num_pages = len(pdf_reader.pages)
            text = []
            for page in range(num_pages):
                text.append(pdf_reader.pages[page].extract_text())
        return "".join(text)

    @staticmethod
    def write_pdf(file_path, text):
        pdf_writer = PyPDF2.PdfWriter()

        # 创建一个新的PDF页面并设置文本内容
        page = PyPDF2.PageObject()
        page.extract_text = lambda: text

        # 将页面添加到PDF写入器中
        pdf_writer.add_page(page)

        # 将写入器的内容写入到指定的PDF文件中
        with open(file_path, 'wb') as pdf_file:
            pdf_writer.write(pdf_file)

    # @staticmethod
    # def pdf_to_image(pdf_path, image_folder):
    #     """pdf文件转换成image文件"""
    #     pdf_document = fitz.open(pdf_path)
    #     for page_number in range(pdf_document.page_count):
    #         page = pdf_document.load_page(page_number)
    #         image = page.get_pixmap(matrix=fitz.Matrix(300 / 72, 300 / 72))  # 300 DPI
    #         image_path = "%s/page_%s.png" % (image_folder, page_number + 1)
    #         image.save(image_path)
    #     pdf_document.close()

    @staticmethod
    def get_pdf_properties(file_path):
        pdf = open(file_path, 'rb')
        pdf_reader = PyPDF2.PdfReader(pdf)
        properties = pdf_reader.metadata
        modified_time = properties.get('/ModDate')
        created_time = properties.get('/CreationDate')
        pdf.close()

        return created_time, modified_time


class HtmlFileHandleBs4:
    """使用bs4解析html的数据"""

    def __init__(self, html):
        self.soup = BeautifulSoup(html, 'lxml')

    def find_all(self, label: dict):
        """
        返回所有匹配的标签
        name：要查找的标签名，可以是字符串或正则表达式。  name='a'
        attrs：要匹配的属性及其值，以字典形式传入。    attrs={'class': 'title'}
        text：要匹配的文本内容，可以是字符串或正则表达式。 text='Hello'
        limit：限制返回结果的数量。
        """
        tags = self.soup.find_all(**label)
        return tags

    def find(self, label: dict):
        """只返回第一个匹配的标签"""
        return self.soup.find(**label)

    def select(self, css_label: dict):
        """
        使用 CSS 选择器语法查找并返回匹配的标签列表
        css_label = {"selector": ".title"} 查找 class 属性为 'title' 的标签
        """
        tags = self.soup.select(**css_label)
        return tags

    @staticmethod
    def get_parent(tag):
        """parent 属性：获取当前标签的父标签"""
        return tag.parent

    @staticmethod
    def get_contents(tag):
        """contents 属性：获取当前标签的所有子标签列表"""
        child_tags = tag.contents
        return child_tags

    @staticmethod
    def get_string(tag):
        """string 属性：获取当前标签的文本内容"""
        return tag.string

    @staticmethod
    def get_attr(tag, attr):
        """获取标签属性值属性值"""
        return tag.get(attr)  # tag.get("href") 获取当前标签的 href 属性值


class HtmlFileHandleEtree:
    """使用etree解析html的数据"""

    def __init__(self, html):
        self.tree = etree.HTML(html)

    def parse(self, xpath):
        """注意元素下标是从1开始的"""
        return self.tree.xpath(xpath)


class JavaScriptFileHandle:
    """python 执行js代码 需要本地有 node.js 环境"""

    @staticmethod
    def read_js(file_path):
        with open(file_path, "r", encoding="utf-8") as f:
            js_code = f.read()

        return js_code

    @staticmethod
    def run_js_code(js_code, args: tuple):
        """
        args 元组传参 第一个参数为js函数名称 后面的参数是js函数的参数
        """
        context = execjs.compile(js_code)
        return context.call(*args)


class TemplateLoader:
    def __init__(self, template_folder):
        self.env = Environment(loader=FileSystemLoader(template_folder))

    def load_template(self, template_name, kwargs: dict):
        template = self.env.get_template(template_name)
        return template.render(**kwargs)


class ImageFileHandle:
    @staticmethod
    def check_image_file(file_path):
        return imghdr.what(file_path)

    @staticmethod
    def to_grey_image(image=None, file_path=None):
        """将彩色图像转换成灰度图像 便于ocr识别"""
        if file_path:
            image = cv2.imread(file_path)
        return cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)


class VideoFileHandle:
    def __init__(self, file_path):
        self.video = VideoFileClip(file_path)

    def get_video_frames(self):
        return self.video.iter_frames()


class GuiFileHandle:
    def __init__(self):
        """若已经实例化了tkinter，可以不实例化此类"""
        self.root = tk.Tk()
        self.root.withdraw()

    @staticmethod
    def select_file_path(file_type):
        file_path = filedialog.askopenfilename(filetypes=[(file_type, "*." + file_type)])
        return [file_path]

    @staticmethod
    def select_folder_path(need_file=True, file_type=None):
        folder_path = filedialog.askdirectory()
        if need_file:
            return OtherFileHandle.get_all_file_in_folder(folder_path, file_type)
        return folder_path


class OtherFileHandle:
    @staticmethod
    def create_file(file_path):
        if not os.path.exists(file_path):
            if str(file_path).endswith("xlsx"):
                workbook = openpyxl.Workbook()
                workbook.save(filename=file_path)
            else:
                open(file_path, 'w').close()
            logger.info("create file and filepath is %s" % file_path)

    @staticmethod
    def create_folder(folder_path):
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
            logger.info("create folder and folder path is %s" % folder_path)

    @staticmethod
    def remove_source(file_path):
        if os.path.exists(file_path):
            os.remove(file_path)

    @staticmethod
    def get_same_level_file_path(file_path=None):
        """获取同级目录文件 不能跨目录调用"""
        if not file_path:
            return os.path.dirname(os.path.abspath(__file__))
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), file_path)

    @staticmethod
    def get_all_file_in_folder(folder_path, file_type=None):
        """以~开头的文件一般表示临时文件"""
        if file_type:
            files = [file for file in os.listdir(folder_path) if file.endswith("." + file_type)]
        else:
            files = [file for file in os.listdir(folder_path)]
        file_path_list = [os.path.join(folder_path, file) for file in files if "~" not in file]
        return file_path_list


if __name__ == '__main__':
    # for a in XlsxFileHandle.read_xlsx(r"D:\Desktop\9.8 版纳-昆明.xlsx"):
    #     logger.info(a)

    # OtherHandle.pdf_to_image(r"D:\Desktop\irm\2020-02-03-1207288434.pdf", r"D:\Desktop\test")
    # logger.info(OtherHandle.get_all_file_in_folder(r"D:\Desktop\siba", file_type="sh"))
    # XlsxFileHandle.write_xlsx(r"D:\Desktop\test.xlsx", dataframe=[["cj", "17"]], columns_name=["name", "age"])
    # XlsxFileHandle.write_xlsx(r"D:\Desktop\test.xlsx", dataframe=[["tsl", "18"]], is_add=True)

    # GuiFileHandle().select_folder_path()
    # logger.info(OtherHandle.get_same_level_file_path("SRConfig.py"))

    # js = """
    #         function add(num1, num2) {
    #         return num1 + num2;
    #     }
    # """
    # logger.info(JavaScriptFileHandle.run_js_code(js, ("add", 1, 2)))

    # logger.info(PdfFileHandle.read_pdf(r"D:\Desktop\Temp\乌鲁木齐地窝堡.pdf"))

    logger.info(PdfFileHandle.get_pdf_properties(r"D:\Desktop\CdsDocuments\cloud_os_demand\cloud_os接口设计.pdf"))
